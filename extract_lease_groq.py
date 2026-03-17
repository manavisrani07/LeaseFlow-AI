from __future__ import annotations

import argparse
import json
import os
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any
from urllib.parse import urlparse
from uuid import uuid4

from docx import Document
from docxtpl import DocxTemplate
from dotenv import load_dotenv
from groq import Groq
from pypdf import PdfReader
from supabase import Client, create_client

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
MATERIALS_DIR = BASE_DIR / "materials"
OUTPUT_DIR = MATERIALS_DIR / "generated_welcome_packs"
TEMPLATE_PATH = MATERIALS_DIR / "Tenant Welcome Pack Template.docx"
DEFAULT_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"
DEFAULT_BUCKET = "leases"
DEFAULT_TABLE = "leases"

SCHEMA_KEYS = [
    "tenant_name",
    "property_address",
    "lease_start_date",
    "lease_end_date",
    "rent_amount",
    "bond_amount",
    "num_occupants",
    "pet_permission",
    "parking_included",
    "special_conditions",
    "landlord_name",
    "property_manager_name",
    "property_manager_email",
    "property_manager_phone",
]


@dataclass
class LeaseProcessingResult:
    lease_id: str
    filename: str
    status: str
    extracted: dict[str, str]
    model_response: str
    storage_path_original: str
    storage_path_generated: str


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def sanitize_filename(name: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9._ -]+", "_", name).strip()
    return safe.replace(" ", "_") or "lease"


def read_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def read_pdf_text(pdf_path: Path) -> str:
    reader = PdfReader(str(pdf_path))
    pages: list[str] = []

    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(text)

    return "\n\n".join(pages).strip()


def read_document_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return read_docx_text(file_path)
    if suffix == ".pdf":
        return read_pdf_text(file_path)
    raise ValueError(f"Unsupported file type: {file_path.suffix}")


def build_prompt(doc_text: str) -> str:
    schema_template = {key: "" for key in SCHEMA_KEYS}

    return (
        "You are an expert lease information extraction assistant. "
        "Extract details from lease text. Return JSON only, no markdown, no explanation. "
        "Use the exact keys shown in the schema and return string values. "
        "If missing, return an empty string.\n\n"
        "Notes for extraction:\n"
        "- tenant_name: include all tenant names as a single comma-separated string if multiple.\n"
        "- rent_amount and bond_amount: keep currency and amount exactly as written, if rent_amount is given fortnightly change it to monthly means multiply by 2. and just mention AUD at the end of the amount no other text.\n"
        "- pet_permission: 'Not permitted' or a description including any conditions\n"
        "- parking_included: 'Not included' or details such as space number and location\n"
        "- special_conditions: summarize special clauses briefly if present, else empty string.\n\n"
        f"Required JSON schema example:\n{json.dumps(schema_template, indent=2)}\n\n"
        "LEASE TEXT:\n"
        f"{doc_text}"
    )


def normalize_result(result: dict[str, Any]) -> dict[str, str]:
    normalized: dict[str, str] = {}
    for key in SCHEMA_KEYS:
        value = result.get(key)
        if value is None:
            normalized[key] = ""
            continue

        if isinstance(value, list):
            normalized[key] = ", ".join(str(item).strip() for item in value if str(item).strip())
            continue

        normalized[key] = str(value).strip()

    return normalized


def parse_model_json(content: str) -> dict[str, str]:
    try:
        parsed = json.loads(content)
    except json.JSONDecodeError:
        start = content.find("{")
        end = content.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise ValueError("Model response is not valid JSON")
        parsed = json.loads(content[start : end + 1])

    if not isinstance(parsed, dict):
        raise ValueError("Model response JSON must be an object")

    return normalize_result(parsed)


def call_groq_for_extraction(client: Groq, model: str, doc_text: str) -> tuple[dict[str, str], str]:
    prompt = build_prompt(doc_text)

    completion = client.chat.completions.create(
        model=model,
        temperature=0,
        max_completion_tokens=1200,
        messages=[{"role": "user", "content": prompt}],
    )

    content = completion.choices[0].message.content or "{}"
    normalized = parse_model_json(content)
    return normalized, content


def generate_welcome_pack(data: dict[str, str], output_path: Path, template_path: Path = TEMPLATE_PATH) -> Path:
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = DocxTemplate(str(template_path))
    context = {key: data.get(key, "") for key in SCHEMA_KEYS}
    doc.render(context)
    doc.save(str(output_path))

    special_conditions = (data.get("special_conditions") or "").strip()
    if not special_conditions:
        rendered_doc = Document(str(output_path))
        paragraphs = rendered_doc.paragraphs

        for idx, paragraph in enumerate(paragraphs):
            if paragraph.text.strip().lower() == "special conditions":
                heading_element = paragraph._element
                heading_element.getparent().remove(heading_element)

                if idx < len(paragraphs) - 1:
                    next_paragraph = paragraphs[idx + 1]
                    next_text = next_paragraph.text.strip()
                    if not next_text or next_text == "{{special_conditions}}":
                        next_element = next_paragraph._element
                        next_element.getparent().remove(next_element)
                break

        rendered_doc.save(str(output_path))

    return output_path


class LeaseProcessor:
    def __init__(
        self,
        groq_api_key: str,
        supabase_url: str,
        supabase_service_key: str,
        model: str = DEFAULT_MODEL,
        bucket: str = DEFAULT_BUCKET,
        table: str = DEFAULT_TABLE,
    ) -> None:
        self.model = model
        self.bucket = bucket
        self.table = table
        self.groq_client = Groq(api_key=groq_api_key)
        self.supabase: Client = create_client(supabase_url, supabase_service_key)

    def _insert_processing_row(
        self,
        lease_id: str,
        filename: str,
        storage_path_original: str,
        storage_path_generated: str,
    ) -> None:
        payload = {
            "id": lease_id,
            "filename": filename,
            "storage_path_original": storage_path_original,
            "storage_path_generated": storage_path_generated,
            "status": "processing",
            "created_at": utc_now_iso(),
            "updated_at": utc_now_iso(),
        }
        self.supabase.table(self.table).insert(payload).execute()

    def _update_row(self, lease_id: str, payload: dict[str, Any]) -> None:
        payload["updated_at"] = utc_now_iso()
        self.supabase.table(self.table).update(payload).eq("id", lease_id).execute()

    def _upload_to_storage(self, storage_path: str, file_bytes: bytes, content_type: str) -> None:
        self.supabase.storage.from_(self.bucket).upload(
            path=storage_path,
            file=file_bytes,
            file_options={"content-type": content_type, "upsert": "false"},
        )

    def get_lease_row(self, lease_id: str) -> dict[str, Any] | None:
        response = (
            self.supabase.table(self.table)
            .select("id, filename, status, extracted, storage_path_generated, storage_path_original, updated_at")
            .eq("id", lease_id)
            .limit(1)
            .execute()
        )
        rows = response.data or []
        if not rows:
            return None
        row = rows[0]
        if not isinstance(row, dict):
            return None
        return row

    def download_generated_welcome_pack(self, lease_id: str) -> tuple[str, bytes]:
        row = self.get_lease_row(lease_id)
        if not row:
            raise ValueError("Lease not found")

        storage_path = row.get("storage_path_generated")
        if not storage_path:
            raise ValueError("Generated welcome pack is not available")

        file_bytes = self.supabase.storage.from_(self.bucket).download(storage_path)
        if not file_bytes:
            raise ValueError("Failed to download generated welcome pack")

        filename = sanitize_filename(Path(str(storage_path)).name)
        return filename, file_bytes

    def list_leases(self, limit: int = 100) -> list[dict[str, Any]]:
        response = (
            self.supabase.table(self.table)
            .select(
                "id, filename, status, extracted, storage_path_generated, storage_path_original, created_at, updated_at"
            )
            .order("created_at", desc=True)
            .limit(limit)
            .execute()
        )
        rows = response.data or []
        return [row for row in rows if isinstance(row, dict)]

    def runtime_config(self) -> dict[str, str]:
        raw_url = getattr(self.supabase, "supabase_url", "")
        host = urlparse(str(raw_url)).netloc or "unknown"
        return {
            "model": self.model,
            "storage_bucket": self.bucket,
            "table": self.table,
            "supabase_host": host,
        }

    def process_file(self, file_path: Path, filename: str | None = None) -> LeaseProcessingResult:
        display_name = filename or file_path.name
        safe_name = sanitize_filename(display_name)
        safe_stem = sanitize_filename(Path(display_name).stem)
        suffix = file_path.suffix.lower()

        if suffix not in {".pdf", ".docx"}:
            raise ValueError("Only PDF and DOCX files are supported")

        lease_id = str(uuid4())
        original_storage_path = f"originals/{lease_id}/{safe_name}"
        generated_storage_path = f"generated/{lease_id}/{safe_stem}_welcome_pack.docx"

        self._insert_processing_row(
            lease_id=lease_id,
            filename=display_name,
            storage_path_original=original_storage_path,
            storage_path_generated=generated_storage_path,
        )

        try:
            original_bytes = file_path.read_bytes()
            original_content_type = (
                "application/pdf"
                if suffix == ".pdf"
                else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            self._upload_to_storage(original_storage_path, original_bytes, original_content_type)

            doc_text = read_document_text(file_path)
            if not doc_text:
                raise ValueError("No text could be extracted from file")

            extracted, model_response = call_groq_for_extraction(self.groq_client, self.model, doc_text)

            with TemporaryDirectory() as tmp_dir:
                generated_path = Path(tmp_dir) / f"{safe_stem}_welcome_pack.docx"
                generate_welcome_pack(extracted, generated_path)
                self._upload_to_storage(
                    generated_storage_path,
                    generated_path.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            update_payload = {
                "storage_path_original": original_storage_path,
                "storage_path_generated": generated_storage_path,
                "status": "completed",
                "extracted": extracted,
                "model_response": model_response,
            }
            self._update_row(lease_id, update_payload)

            return LeaseProcessingResult(
                lease_id=lease_id,
                filename=display_name,
                status="completed",
                extracted=extracted,
                model_response=model_response,
                storage_path_original=original_storage_path,
                storage_path_generated=generated_storage_path,
            )
        except Exception as exc:
            self._update_row(
                lease_id,
                {
                    "status": "failed",
                    "storage_path_original": original_storage_path,
                    "storage_path_generated": generated_storage_path,
                    "model_response": str(exc),
                },
            )
            raise


def resolve_required_env() -> tuple[str, str, str, str]:
    groq_api_key = os.getenv("GROQ_API_KEY")
    supabase_url = os.getenv("SUPBASE_URL")
    supabase_service_key = os.getenv("SUPBASE_SERVICE_KEY")
    model = os.getenv("GROQ_MODEL", DEFAULT_MODEL)

    missing = []
    if not groq_api_key:
        missing.append("GROQ_API_KEY")
    if not supabase_url:
        missing.append("SUPBASE_URL")
    if not supabase_service_key:
        missing.append("SUPBASE_SERVICE_KEY")

    if missing:
        raise EnvironmentError(f"Missing environment variables: {', '.join(missing)}")

    return groq_api_key, supabase_url, supabase_service_key, model


def build_processor_from_env() -> LeaseProcessor:
    groq_api_key, supabase_url, supabase_service_key, model = resolve_required_env()
    return LeaseProcessor(
        groq_api_key=groq_api_key,
        supabase_url=supabase_url,
        supabase_service_key=supabase_service_key,
        model=model,
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Process a lease file and persist extraction plus generated welcome pack in Supabase."
    )
    parser.add_argument(
        "--file",
        type=str,
        required=True,
        help="Lease file path (.docx or .pdf)",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    file_path = Path(args.file).resolve()
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    processor = build_processor_from_env()
    result = processor.process_file(file_path)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_json = OUTPUT_DIR / f"{sanitize_filename(file_path.stem)}_extracted.json"
    output_json.write_text(json.dumps(result.extracted, indent=2, ensure_ascii=False), encoding="utf-8")

    print(f"Lease ID: {result.lease_id}")
    print(f"Status: {result.status}")
    print(f"Original storage path: {result.storage_path_original}")
    print(f"Generated storage path: {result.storage_path_generated}")
    print(f"Saved local extraction JSON: {output_json}")


if __name__ == "__main__":
    main()
